#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract text and images from a .docx file."""
import os
import re
import sys
import zipfile
from docx import Document

DOCX_PATH = r"d:\6\恋爱小程序\反馈.docx"
OUT_DIR = r"d:\6\恋爱小程序\反馈媒体"
MD_PATH = r"d:\6\恋爱小程序\反馈.md"

os.makedirs(OUT_DIR, exist_ok=True)

doc = Document(DOCX_PATH)

# Extract text with basic style hints
lines = []
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name if para.style else "Normal"
    if style.startswith("Heading"):
        level = re.sub(r"\D", "", style) or "2"
        lines.append(f"{'#' * int(level)} {text}")
    elif style.startswith("List") or style.startswith("Bullet"):
        lines.append(f"- {text}")
    else:
        lines.append(text)

# Extract tables
for table_idx, table in enumerate(doc.tables, 1):
    lines.append(f"\n[Table {table_idx}]")
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(" | ".join(cells))

# Extract images by unzipping media
image_refs = []
with zipfile.ZipFile(DOCX_PATH, "r") as z:
    media_files = [n for n in z.namelist() if n.startswith("word/media/")]
    for idx, name in enumerate(sorted(media_files), 1):
        ext = os.path.splitext(name)[1]
        out_name = f"image{idx:03d}{ext}"
        out_path = os.path.join(OUT_DIR, out_name)
        with z.open(name) as src, open(out_path, "wb") as dst:
            dst.write(src.read())
        image_refs.append(f"![image{idx:03d}](反馈媒体/{out_name})")

with open(MD_PATH, "w", encoding="utf-8") as f:
    f.write("# 反馈文档提取\n\n")
    f.write("## 文本内容\n\n")
    f.write("\n\n".join(lines))
    f.write("\n\n## 图片内容\n\n")
    f.write("\n\n".join(image_refs))
    f.write("\n")

print(f"Text extracted to: {MD_PATH}")
print(f"Images extracted to: {OUT_DIR}")
print(f"Image count: {len(image_refs)}")
